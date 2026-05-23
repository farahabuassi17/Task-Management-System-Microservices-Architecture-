from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Task_Management_System_Final_Documentation.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True)
        set_cell_shading(hdr_cells[i], "E9EEF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    document.add_paragraph()
    return table


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def style_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "Task Management System - Final Project Documentation"
    header_p.runs[0].font.name = "Arial"
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_section_title(document, title):
    document.add_heading(title, level=1)


def add_subtitle(document, title):
    document.add_heading(title, level=2)


def build_document():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Task Management System")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Final Project Documentation")
    run.font.name = "Arial"
    run.font.size = Pt(15)
    run.bold = True

    meta = [
        ("Student Name", "Farah M T AbuAssi"),
        ("Student ID", "220221408"),
        ("Supervisor", "Dr. Abdelkareem Alashqar"),
        ("Course", "Advanced Software Engineering"),
        ("Document Date", "May 20, 2026"),
    ]
    add_table(doc, ["Field", "Value"], meta, widths=[2.2, 4.8])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "This document consolidates the documented design, communication styles, "
        "implementation details, containerization work, and CI/CD preparation completed "
        "throughout the semester."
    )
    doc.add_page_break()

    add_section_title(doc, "Table of Contents")
    toc_items = [
        "1. Project Overview",
        "2. Domain Model and Bounded Contexts",
        "3. Microservices Architecture",
        "4. Communication Styles and Technologies",
        "5. API Contracts and Message Schemas",
        "6. Database Ownership",
        "7. Implementation Summary",
        "8. Running and Manual Testing",
        "9. Docker Containerization",
        "10. GitHub Monorepo and CI/CD",
        "11. Quality Review and Final Status",
        "12. Conclusion",
    ]
    add_bullets(doc, toc_items)
    doc.add_page_break()

    add_section_title(doc, "1. Project Overview")
    doc.add_paragraph(
        "The Task Management System is a microservices-based application designed to manage "
        "users, tasks, task assignment, task status changes, and notifications. Unlike a "
        "monolithic system, the application is decomposed into independently deployable services. "
        "Each service is responsible for a specific business capability and owns its private data store."
    )
    doc.add_paragraph(
        "The project applies Domain-Driven Design principles by defining bounded contexts, "
        "protecting internal models, and exposing only explicit contracts between services. "
        "The final implementation includes the main required services: User Service, Task Service, "
        "and Notification Service. It also includes supporting services and delivery assets such as "
        "an API Gateway, Board Service, Dockerfile, GitHub Actions workflow, and formal project documentation."
    )

    add_section_title(doc, "2. Domain Model and Bounded Contexts")
    add_subtitle(doc, "2.1 Identity and Access Context")
    doc.add_paragraph(
        "The Identity and Access Context is responsible for managing user profiles and user identity data. "
        "It owns the User aggregate and hides internal details such as authentication logs or sensitive "
        "identity implementation details. The Task Service depends only on a shared user summary model."
    )
    add_subtitle(doc, "2.2 Task Board Context")
    doc.add_paragraph(
        "The Task Board Context is the core business context. It manages tasks, task status transitions, "
        "priorities, descriptions, board identifiers, and assignments. It owns the Task aggregate and "
        "exposes task operations through REST and GraphQL APIs."
    )
    add_subtitle(doc, "2.3 Notification Context")
    doc.add_paragraph(
        "The Notification Context is responsible for receiving task-related events and storing notification "
        "records. It is autonomous and does not need direct knowledge of Task Service internals. This supports "
        "loose coupling and independent deployability."
    )

    add_table(
        doc,
        ["Bounded Context", "Microservice", "Owned Data", "Main Responsibility"],
        [
            ("Identity and Access", "User Service", "User DB", "Manage users and validate active assignees."),
            ("Task Board", "Task Service", "Task DB", "Create, retrieve, and update tasks."),
            ("Notification", "Notification Service", "Notification DB", "Consume task events and store notifications."),
            ("Board Support", "Board Service", "Board DB", "Handle asynchronous board summary requests."),
        ],
        widths=[1.7, 1.7, 1.3, 3.0],
    )

    add_section_title(doc, "3. Microservices Architecture")
    doc.add_paragraph(
        "The architecture is organized as a public GitHub monorepo. Each service has its own source code, "
        "configuration, Maven build file, and database configuration. This structure simplifies management "
        "while preserving service independence."
    )
    add_table(
        doc,
        ["Service", "Port", "Technology", "Purpose"],
        [
            ("user-service", "8081 / 9091", "Spring Boot, REST, gRPC, MySQL", "User management and gRPC validation."),
            ("task-service", "8082", "Spring Boot, REST, GraphQL, gRPC client, RabbitMQ, MySQL", "Task lifecycle and event publishing."),
            ("notification-service", "8083", "Spring Boot, RabbitMQ, REST, MySQL", "Notification event consumption and storage."),
            ("api-gateway", "8080", "Spring Boot REST proxy", "Single entry point for external clients."),
            ("board-service", "8084", "Spring Boot, RabbitMQ, REST, MySQL", "Board summary request handling."),
        ],
        widths=[1.5, 1.0, 2.5, 2.8],
    )

    add_section_title(doc, "4. Communication Styles and Technologies")
    doc.add_paragraph(
        "The system uses multiple communication styles selected according to the needs of each interaction."
    )
    add_table(
        doc,
        ["Interaction", "Style", "Technology", "Reason"],
        [
            (
                "External client to services",
                "Synchronous request-response",
                "REST over HTTP using JSON",
                "Simple, human-readable, and widely supported by Spring Boot.",
            ),
            (
                "Task Service to User Service",
                "Synchronous request-response",
                "gRPC with Protocol Buffers",
                "Efficient internal validation before saving tasks.",
            ),
            (
                "Task Service to Notification Service",
                "Asynchronous event-driven",
                "RabbitMQ with JSON payloads",
                "Keeps notification processing independent from task creation.",
            ),
            (
                "Task Service external querying",
                "Flexible client-driven query",
                "GraphQL",
                "Allows clients to request only the fields they need.",
            ),
            (
                "Task Service to Board Service",
                "Asynchronous request flow",
                "RabbitMQ with JSON payloads",
                "Supports long-running board summary work without blocking clients.",
            ),
        ],
        widths=[1.7, 1.4, 1.8, 2.8],
    )

    add_section_title(doc, "5. API Contracts and Message Schemas")
    add_subtitle(doc, "5.1 User Service REST Contract")
    doc.add_paragraph("Endpoint: GET /api/users/{id}")
    doc.add_paragraph(
        'Example response: {"id":"U100","fullName":"Farah AbuAssi","email":"farah@example.com","role":"PROJECT_MANAGER","active":true}'
    )

    add_subtitle(doc, "5.2 Task Service REST Contract")
    doc.add_paragraph("Endpoint: POST /api/tasks")
    doc.add_paragraph(
        'Example request: {"title":"Complete Assignment","description":"Finish implementation","priority":"HIGH","boardId":"B-001","assignedUserId":"U100"}'
    )
    doc.add_paragraph("Endpoint: PATCH /api/tasks/{id}/status")
    doc.add_paragraph('Example request: {"status":"IN_PROGRESS"}')

    add_subtitle(doc, "5.3 Notification Event Contract")
    doc.add_paragraph("Exchange: task.events")
    doc.add_paragraph("Routing key: task.events.updates")
    doc.add_paragraph(
        'Event payload includes eventId, eventType, timestamp, taskId, taskTitle, oldStatus, newStatus, assigneeEmail, and alertMessage.'
    )

    add_subtitle(doc, "5.4 GraphQL Contract")
    doc.add_paragraph(
        "The Task Service exposes GraphQL at /graphql and GraphiQL at /graphiql. It supports task queries, "
        "task creation, and task status updates."
    )

    add_section_title(doc, "6. Database Ownership")
    doc.add_paragraph(
        "Each microservice owns a separate MySQL database. This avoids common coupling and allows each service "
        "to evolve independently."
    )
    add_table(
        doc,
        ["Service", "Database", "Main Tables"],
        [
            ("User Service", "user_service_db", "users"),
            ("Task Service", "task_service_db", "tasks"),
            ("Notification Service", "notification_service_db", "notifications"),
            ("Board Service", "board_service_db", "board_summaries"),
        ],
        widths=[2.2, 2.4, 2.6],
    )

    add_section_title(doc, "7. Implementation Summary")
    add_bullets(
        doc,
        [
            "User Service was implemented as the source of truth for user profile data.",
            "Task Service was implemented with REST, GraphQL, gRPC client logic, and RabbitMQ publishing.",
            "Notification Service was implemented as an independent RabbitMQ consumer with its own database.",
            "Task Service validates users through gRPC before saving tasks.",
            "Task Service publishes task status events after create and update operations.",
            "Task event publishing was made resilient so task creation does not fail if RabbitMQ is temporarily unavailable.",
            "Dockerfile was added for Task Service containerization.",
            "GitHub Actions workflow was added for automated build and Docker image publishing.",
        ],
    )

    add_section_title(doc, "8. Running and Manual Testing")
    doc.add_paragraph("Required local dependencies:")
    add_bullets(doc, ["MySQL on localhost:3306", "RabbitMQ on localhost:5672", "Java 17", "Maven or Maven Wrapper"])
    add_table(
        doc,
        ["Step", "Command or URL"],
        [
            ("Run User Service", "cd user-service && mvn spring-boot:run"),
            ("Run Notification Service", "cd notification-service && .\\mvnw.cmd spring-boot:run"),
            ("Run Task Service", "cd task-service && mvn spring-boot:run"),
            ("Test user", "GET http://localhost:8081/api/users/U100"),
            ("Create task", "POST http://localhost:8082/api/tasks"),
            ("Update status", "PATCH http://localhost:8082/api/tasks/{id}/status"),
            ("View notifications", "GET http://localhost:8083/api/notifications"),
            ("Open GraphiQL", "http://localhost:8082/graphiql"),
        ],
        widths=[2.2, 5.2],
    )

    add_section_title(doc, "9. Docker Containerization")
    doc.add_paragraph(
        "At least one microservice was prepared for Docker deployment. The Task Service contains a Dockerfile "
        "that builds the application JAR and runs it in an isolated Java runtime container."
    )
    doc.add_paragraph("Docker build command:")
    doc.add_paragraph("docker build -t task-service:latest .")
    doc.add_paragraph("Docker run command:")
    doc.add_paragraph("docker run --rm -p 8082:8082 task-service:latest")

    add_section_title(doc, "10. GitHub Monorepo and CI/CD")
    doc.add_paragraph(
        "The project is structured as a monorepo. A single public GitHub repository should contain all service "
        "directories and project documentation."
    )
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ("user-service/", "User Service source code"),
            ("task-service/", "Task Service source code and Dockerfile"),
            ("notification-service/", "Notification Service source code"),
            ("api-gateway/", "Gateway service"),
            ("board-service/", "Board support service"),
            (".github/workflows/ci-cd.yml", "GitHub Actions CI/CD workflow"),
            ("README.md", "Repository overview and usage instructions"),
            ("PROJECT_INFO.md", "Project information summary"),
        ],
        widths=[2.8, 4.4],
    )
    doc.add_paragraph(
        "The CI/CD workflow builds the Task Service JAR, builds a Docker image, and pushes it to Docker Hub. "
        "The workflow requires the GitHub repository secrets DOCKERHUB_USERNAME and DOCKERHUB_TOKEN."
    )

    add_section_title(doc, "11. Quality Review and Final Status")
    add_table(
        doc,
        ["Requirement", "Status", "Evidence"],
        [
            ("Complete all designed microservices", "Completed", "User Service, Task Service, and Notification Service are implemented."),
            ("gRPC validation", "Completed", "Task Service calls User Service before saving a task."),
            ("Independent MySQL databases", "Completed", "Each service has its own database configuration."),
            ("Public GitHub Monorepo", "Prepared", "Monorepo structure is ready; repository must be pushed publicly."),
            ("Docker containerization", "Prepared", "Task Service includes Dockerfile."),
            ("CI/CD bonus", "Prepared", "GitHub Actions workflow is included."),
            ("Manual testing", "Partially completed", "Compile succeeds; runtime testing requires MySQL and RabbitMQ."),
        ],
        widths=[2.4, 1.4, 3.6],
    )

    add_section_title(doc, "12. Conclusion")
    doc.add_paragraph(
        "The Task Management System demonstrates a complete microservices implementation based on the documented "
        "design. The system separates responsibilities into independently deployable services, uses synchronous "
        "gRPC validation where immediate confirmation is required, and uses asynchronous RabbitMQ events for "
        "notification processing. The project also includes Docker and CI/CD preparation, making it suitable for "
        "formal submission and future deployment improvements."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
